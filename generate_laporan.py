from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(13, 110, 253))
    return p

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(0, 70, 180))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(30, 30, 30))
    p.paragraph_format.space_before = Pt(10)
    return p

def body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.space_after = Pt(4)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F4FF')
    p._p.get_or_add_pPr().append(shading)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

# ===================== COVER =====================
heading1("LAPORAN ANALISIS KODE PROGRAM")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Pemrograman Web 1 — Sesi 8 UTS")
set_font(r, size=12, italic=True, color=(80,80,80))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Muhammad Dava Fahreza")
set_font(r2, size=12, bold=True)

doc.add_paragraph()

# ===================== JUDUL SECTION =====================
heading2("A. SOURCE CODE APLIKASI")
body("Project ini terdiri dari dua halaman HTML utama:")
body("  1. index.html  — Halaman utama portfolio")
body("  2. contact.html — Halaman kontak")
body("Didukung oleh dua file CSS: style.css dan contact.css, serta data TypeScript di folder data/.")
doc.add_paragraph()

heading2("C. ANALISIS SEMUA TAG HTML")
body("Berikut analisis elemen HTML pertama yang muncul/satu elemen terkait dari project ini.")
doc.add_paragraph()

tags = [
    ("1. <!DOCTYPE html>",
     "Deklarasi tipe dokumen. Memberitahu browser bahwa dokumen menggunakan standar HTML5.",
     "index.html — baris 1:",
     "<!doctype html>"),
    ("2. <html>",
     "Elemen root/akar dari seluruh halaman HTML. Atribut lang='en' menentukan bahasa dokumen untuk aksesibilitas dan SEO.",
     "index.html — baris 2:",
     '<html lang="en">'),
    ("3. <head>",
     "Wadah untuk metadata halaman — tidak ditampilkan ke pengguna tetapi penting untuk browser, SEO, dan resource loading.",
     "index.html — baris 4:",
     "<head>\n    <meta charset=\"utf-8\">\n    <title>Muhammad Dava Fahreza</title>\n</head>"),
    ("4. <meta>",
     "Tag metadata tanpa penutup. charset mengatur encoding karakter, viewport mengatur tampilan responsif di perangkat mobile.",
     "index.html — baris 5-6:",
     '<meta charset="utf-8">\n<meta name="viewport" content="width=device-width, initial-scale=1">'),
    ("5. <title>",
     "Menentukan judul halaman yang tampil di tab browser dan hasil mesin pencari.",
     "index.html — baris 7:",
     "<title>Muhammad Dava Fahreza</title>"),
    ("6. <link>",
     "Menghubungkan dokumen eksternal seperti file CSS dan icon. Tag tanpa penutup.",
     "index.html — baris 9-14:",
     '<link rel="stylesheet" href="style.css">\n<link rel="stylesheet" href="contact.css">\n<link rel="icon" type="image/png" href="favicon/logo-mdf-1.jpeg" />'),
    ("7. <body>",
     "Wadah untuk semua konten yang ditampilkan di halaman — teks, gambar, tombol, tabel, dan elemen lainnya.",
     "index.html — baris 22:",
     "<body>\n    ...\n</body>"),
    ("8. <nav>",
     "Elemen semantik untuk navigasi utama halaman. Membantu aksesibilitas dan SEO.",
     "index.html — baris 24:",
     '<nav class="navbar navbar-expand-lg bg-light fixed-top">\n    ...\n</nav>'),
    ("9. <div>",
     "Elemen pembungkus blok generik untuk mengelompokkan elemen dan menerapkan CSS/layout.",
     "index.html — baris 25:",
     '<div class="container-fluid">\n    ...\n</div>'),
    ("10. <button>",
     "Membuat tombol yang bisa diklik. Digunakan untuk toggle navbar mobile dan navigasi sertifikat.",
     "index.html — baris 26:",
     '<button class="navbar-toggler" type="button" data-bs-toggle="collapse">\n    <span class="navbar-toggler-icon"></span>\n</button>'),
]

for tag, desc, impl_label, impl_code in tags:
    heading3(tag)
    body(f"Fungsi: {desc}")
    body(f"Implementasi — {impl_label}")
    code_block(impl_code)
    doc.add_paragraph()

tags2 = [
    ("11. <a>",
     "Hyperlink untuk navigasi antar halaman, link eksternal, dan anchor. Atribut href menentukan tujuan, target='_blank' membuka di tab baru.",
     "index.html — navbar brand dan link GitHub:",
     '<a class="navbar-brand" href="#">MUHAMMAD DAVA FAHREZA</a>\n<a href="http://github.com/mdf05" target="_blank" rel="noopener noreferrer">mdf05</a>'),
    ("12. <ul> dan <li>",
     "<ul> membuat daftar tidak berurutan (unordered list), <li> adalah setiap item dalam daftar. Digunakan untuk menu navbar dan daftar tech stack.",
     "index.html — navbar menu dan tech stack:",
     '<ul class="navbar-nav me-auto">\n    <li class="nav-item"><a class="nav-link" href="#">HOME</a></li>\n    <li class="nav-item"><a class="nav-link" href="contact.html">CONTACT</a></li>\n</ul>\n\n<ul class="tech-list">\n    <li>Next.js</li>\n    <li>React</li>\n</ul>'),
    ("13. <form> dan <input>",
     "<form> membungkus elemen input, <input> adalah field masukan data. Digunakan untuk search bar di navbar.",
     "index.html — search bar:",
     '<form class="d-flex" role="search">\n    <input class="form-control me-2" type="search" placeholder="Search">\n    <button class="btn btn-outline-primary" type="submit">Search</button>\n</form>'),
    ("14. <section>",
     "Elemen semantik yang mendefinisikan bagian/section konten yang tematik dalam halaman.",
     "index.html — section biodata:",
     '<section class="biodata">\n    ...\n</section>'),
    ("15. <h1> – <h5> (Heading)",
     "Tag heading untuk judul dengan tingkat hierarki. h1 paling besar/penting, h5 lebih kecil.",
     "index.html — h2 MY PROFILE dan h5 kategori tech:",
     '<h2 class="text-center my-profile"><b>MY PROFILE</b></h2>\n<h5 class="tech-category-title">🌐 Frontend</h5>'),
    ("16. <img>",
     "Menampilkan gambar. Atribut src menentukan sumber file, alt teks alternatif untuk aksesibilitas.",
     "index.html — foto profil:",
     '<img src="assets/hero/2.png" alt="profile-picture"\n     class="img-thumbnail mx-auto d-block"\n     data-aos="flip-left" data-aos-duration="2000">'),
    ("17. <table>, <thead>, <tbody>, <tr>, <th>, <td>",
     "Elemen tabel HTML. table wadah utama, thead baris header, tbody badan data, tr baris, th sel header (bold), td sel data biasa.",
     "index.html — tabel profil dan tabel pendidikan:",
     '<table class="profile-table">\n    <tr>\n        <th>NAME</th>\n        <td>:</td>\n        <td>MUHAMMAD DAVA FAHREZA</td>\n    </tr>\n</table>\n\n<table class="education-table">\n    <thead><tr><th>Institusi</th><th>Periode</th></tr></thead>\n    <tbody><tr><td>Universitas Siber Asia</td><td>Jan 2026 — Sekarang</td></tr></tbody>\n</table>'),
    ("18. <span>",
     "Elemen inline generik untuk mengelompokkan teks dan menerapkan style tanpa membuat baris baru.",
     "index.html — skill badge:",
     '<span class="skill-badge">FRONTEND</span>\n<span class="skill-badge">BACKEND</span>'),
    ("19. <small>",
     "Menampilkan teks berukuran lebih kecil, umumnya untuk keterangan tambahan atau teks sekunder.",
     "index.html — nama universitas di tabel profil:",
     '<td>INFORMATICS ENGINEERING<br>\n    <small>Universitas Siber Asia (UnSIA)</small>\n</td>'),
    ("20. <br>",
     "Line break, membuat baris baru di dalam elemen teks tanpa membuat paragraf baru. Tag tanpa penutup.",
     "index.html — tabel profil baris STUDY:",
     '<td>INFORMATICS ENGINEERING<br>\n    <small>Universitas Siber Asia (UnSIA)</small>\n</td>'),
    ("21. <p>",
     "Mendefinisikan sebuah paragraf teks. Secara otomatis memberikan margin atas dan bawah.",
     "index.html — paragraf About Me:",
     '<p class="teks-2" data-aos="fade-left" data-aos-duration="3000">\n    Hai, saya <strong>Dava</strong> — mahasiswa ...\n</p>'),
    ("22. <script>",
     "Menyisipkan atau merujuk kode JavaScript. Ditempatkan di akhir body agar tidak memblokir rendering halaman.",
     "index.html — bagian bawah body:",
     '<script src="https://unpkg.com/aos@next/dist/aos.js"></script>\n<script>\n    AOS.init();\n</script>'),
]

for tag, desc, impl_label, impl_code in tags2:
    heading3(tag)
    body(f"Fungsi: {desc}")
    body(f"Implementasi — {impl_label}")
    code_block(impl_code)
    doc.add_paragraph()

# ===================== TEXT FORMATTING TAGS =====================
heading2("TAG TEXT FORMATTING")
body("Berikut tag-tag khusus untuk pemformatan teks yang digunakan dalam project ini.")
doc.add_paragraph()

formatting_tags = [
    ("23. <b> — Bold",
     "Menampilkan teks tebal secara visual, tanpa makna semantik khusus. Berbeda dengan <strong> yang memiliki makna penting secara semantik.",
     "index.html — digunakan di semua judul h2:",
     '<h2 class="text-center my-profile"><b>MY PROFILE</b></h2>\n<h2 class="text-center"><b>ABOUT ME</b></h2>\n<h2 class="text-center education-title"><b>RIWAYAT PENDIDIKAN</b></h2>\n<h2 class="text-center tech-stack-title"><b>TECH STACK</b></h2>\n<h2 class="text-center Sertifikat"><b>SERTIFIKAT SAYA</b></h2>'),
    ("24. <strong> — Strong Importance",
     "Menampilkan teks tebal dengan makna semantik — menandakan teks tersebut penting/krusial. Dibaca screen reader dengan penekanan.",
     "index.html — About Me paragraf (nama, tech stack, status):",
     'Hai, saya <strong>Dava</strong> — mahasiswa ...\n<strong>Next.js</strong>, <strong>React</strong>, <strong>Vue.js</strong>, dan <strong>Laravel</strong>.\nsebagai <strong>freelancer</strong>'),
    ("25. <em> — Emphasis",
     "Menampilkan teks miring dengan makna semantik penekanan. Berbeda dengan <i> yang hanya visual.",
     "index.html — About Me paragraf 2 dan 3:",
     '<em>Nggak cuma frontend</em> — saya juga nyentuh backend ...\n\n<em>Saya tipe orang yang nggak bisa diam,\nselalu penasaran sama hal baru...</em>'),
    ("26. <mark> — Highlight",
     "Menandai/highlight teks yang relevan atau perlu diperhatikan. Default background kuning, di project ini di-override ke biru muda sesuai tema.",
     "index.html — About Me paragraf (keyword penting):",
     'mahasiswa <mark>Teknik Informatika</mark> di Universitas Siber Asia\nbelajar <mark>Machine Learning</mark>\nSejak <mark>Maret 2025</mark> saya aktif kerja sebagai freelancer'),
]

for tag, desc, impl_label, impl_code in formatting_tags:
    heading3(tag)
    body(f"Fungsi: {desc}")
    body(f"Implementasi — {impl_label}")
    code_block(impl_code)
    doc.add_paragraph()

# ===================== TABEL RINGKASAN =====================
heading2("TABEL RINGKASAN TAG HTML")
doc.add_paragraph()

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

header_cells = table.rows[0].cells
headers = ["No", "Tag", "Kategori", "Fungsi Utama"]
for i, h in enumerate(headers):
    p = header_cells[i].paragraphs[0]
    run = p.add_run(h)
    set_font(run, size=10, bold=True, color=(255,255,255))
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), '0D6EFD')
    header_cells[i]._tc.get_or_add_tcPr().append(shading)

rows_data = [
    ("1", "<!DOCTYPE html>", "Deklarasi", "Tipe dokumen HTML5"),
    ("2", "<html>", "Root", "Elemen akar dokumen"),
    ("3", "<head>", "Metadata", "Wadah metadata"),
    ("4", "<meta>", "Metadata", "Charset & viewport"),
    ("5", "<title>", "Metadata", "Judul tab browser"),
    ("6", "<link>", "Resource", "Hubungkan CSS eksternal"),
    ("7", "<body>", "Struktur", "Wadah konten halaman"),
    ("8", "<nav>", "Semantik", "Navigasi halaman"),
    ("9", "<div>", "Layout", "Pembungkus blok"),
    ("10", "<button>", "Interaktif", "Tombol aksi"),
    ("11", "<a>", "Navigasi", "Hyperlink"),
    ("12", "<ul> / <li>", "List", "Daftar tidak berurutan"),
    ("13", "<form> / <input>", "Form", "Input pencarian"),
    ("14", "<section>", "Semantik", "Bagian konten"),
    ("15", "<h1>–<h5>", "Heading", "Judul hierarkis"),
    ("16", "<img>", "Media", "Tampilkan gambar"),
    ("17", "<table>/<tr>/<th>/<td>", "Tabel", "Tampilkan data terstruktur"),
    ("18", "<span>", "Inline", "Wrapper teks inline"),
    ("19", "<small>", "Formatting", "Teks ukuran kecil"),
    ("20", "<br>", "Formatting", "Baris baru"),
    ("21", "<p>", "Teks", "Paragraf"),
    ("22", "<b>", "Text Formatting", "Teks tebal (visual)"),
    ("23", "<strong>", "Text Formatting", "Teks tebal (semantik penting)"),
    ("24", "<em>", "Text Formatting", "Teks miring (semantik penekanan)"),
    ("25", "<mark>", "Text Formatting", "Highlight teks"),
    ("26", "<script>", "JavaScript", "Muat/jalankan JS"),
]

for row_data in rows_data:
    row = table.add_row()
    for i, val in enumerate(row_data):
        p = row.cells[i].paragraphs[0]
        run = p.add_run(val)
        is_formatting = row_data[2] == "Text Formatting"
        set_font(run, size=9, bold=is_formatting, color=(13,110,253) if is_formatting else (30,30,30))

output_path = r"c:\Users\Hype\OneDrive\Desktop\dava\kuliah\Semester-Genap_2025-2026\Pemrograman_Web_1\sesi-8-uts\my--firts-portfolio\Laporan_Analisis_Tag_HTML.docx"
doc.save(output_path)
print(f"File berhasil dibuat: {output_path}")
