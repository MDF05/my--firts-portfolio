from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc_path = r"c:\Users\Hype\OneDrive\Desktop\dava\kuliah\Semester-Genap_2025-2026\Pemrograman_Web_1\sesi-8-uts\my--firts-portfolio\Laporan_Analisis_Tag_HTML.docx"
doc = Document(doc_path)

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(0, 70, 180))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def source_code_line(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(20, 20, 20)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F8F9FF')
    p._p.get_or_add_pPr().append(shading)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    return p

# Page break sebelum source code
doc.add_page_break()

# ===== SOURCE CODE INDEX.HTML =====
heading2("B. SOURCE CODE — index.html")
doc.add_paragraph()

index_html_path = r"c:\Users\Hype\OneDrive\Desktop\dava\kuliah\Semester-Genap_2025-2026\Pemrograman_Web_1\sesi-8-uts\my--firts-portfolio\index.html"
contact_html_path = r"c:\Users\Hype\OneDrive\Desktop\dava\kuliah\Semester-Genap_2025-2026\Pemrograman_Web_1\sesi-8-uts\my--firts-portfolio\contact.html"

with open(index_html_path, 'r', encoding='utf-8') as f:
    index_lines = f.readlines()

for line in index_lines:
    cleaned = line.rstrip('\n')
    source_code_line(cleaned)

# Page break sebelum contact.html
doc.add_page_break()

# ===== SOURCE CODE CONTACT.HTML =====
heading2("B. SOURCE CODE — contact.html")
doc.add_paragraph()

with open(contact_html_path, 'r', encoding='utf-8') as f:
    contact_lines = f.readlines()

for line in contact_lines:
    cleaned = line.rstrip('\n')
    source_code_line(cleaned)

doc.save(doc_path)
print("Source code berhasil ditambahkan ke file Word!")
